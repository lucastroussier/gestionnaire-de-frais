import tkinter as tk
from tkinter import ttk
import openpyxl
from docx import Document
import win32print
import win32api

donnees = []
workbook = openpyxl.load_workbook('sample_data.xlsx')
sheet = workbook.active if workbook.sheetnames else workbook.create_sheet("Frais")



def ouvrir_fenetre_ajout():
    fenetre_ajout.deiconify()

def choisir_type_frais(type):
    global type_de_frais
    type_de_frais = type

def stocker_texte():
    global type_de_frais, donnees, nouvelle_ligne
    texte_entree_frais = entrée_frais.get("1.0", "end-1c")
    texte_entree_date = entre_date.get("1.0", "end-1c")
    nouvelle_ligne = sheet.max_row
    jour, mois, annee = texte_entree_date.split('/')

    if len(jour) == 2 and len(mois) == 2 and len(annee) == 4:  # Vérification des longueurs
        if jour.isdigit() and mois.isdigit() and annee.isdigit():  # Vérification si numérique
            jour = int(jour)
            mois = int(mois)
            annee = int(annee)
            jours_par_mois = [31, 28 if annee % 4 != 0 or (annee % 100 == 0 and annee % 400 != 0) else 29,
                              31, 30, 31, 30, 31, 31, 30, 31, 30, 31]

            if mois >= 1 and mois <= 12 and jour >= 1 and jour <= jours_par_mois[mois - 1]:  # Validation date
                if texte_entree_frais.isdigit():  # Vérification si numérique
                    donnees.append([texte_entree_date, type_de_frais, texte_entree_frais])

                    sheet.cell(row=nouvelle_ligne, column=1, value=texte_entree_date)
                    sheet.cell(row=nouvelle_ligne, column=2, value=type_de_frais)
                    sheet.cell(row=nouvelle_ligne, column=3, value=texte_entree_frais)
                   

                    remboursement = texte_entree_frais
                    if type_de_frais == "hotel" and int(texte_entree_frais) >= 110:
                        remboursement = "110"
                    elif type_de_frais == "repas" and int(texte_entree_frais) >= 30:
                        remboursement = "30"

                    sheet.cell(row=nouvelle_ligne, column=4, value=remboursement)
                    nouvelle_ligne = sheet.max_row + 1  # Déplacer cette ligne ici
                    workbook.save('sample_data.xlsx')

                    tableau_donnee.insert(parent='', index='end', values=(texte_entree_date, type_de_frais, texte_entree_frais, remboursement))
                    entrée_frais.delete("1.0", "end")
                    entre_date.delete("1.0", "end")
                else:
                    print("Le montant doit être un nombre.")
            else:
                print("La date n'est pas valide.")
        else:
            print("La date doit être numérique.")
    else:
        print("Format de date invalide.")

    texte_enregistrement_profil.config(text="enregistrement fini")

def imprimer():
    workbook = openpyxl.load_workbook('sample_data.xlsx')
    sheet = workbook['feuil1'] if 'feuil1' in workbook.sheetnames else workbook.create_sheet("feuil1")
    sheet2 = workbook['feuil2'] if 'feuil2' in workbook.sheetnames else workbook.create_sheet("feuil2")
    nom=sheet2['H2'].value
    prenom=sheet2['I2'].value
    adresse=sheet2['J2'].value
    tel=sheet2['K2'].value
    code_postal=sheet2['L2'].value
    ville=sheet2['M2'].value

    # Charge le classeur Excel
    workbook = openpyxl.load_workbook('sample_data.xlsx')
    sheet = workbook.active if workbook.sheetnames else workbook.create_sheet("Frais")

    # Crée un document Word
    doc = Document("fiche_de_demande_de_remboursement.docx")

    # Récupère les données du classeur Excel
    excel_data = []
    for row in sheet.iter_rows(values_only=True):
        excel_data.append(row)

    # Fusionne les données dans le document Word
    table = doc.add_table(rows=1, cols=len(excel_data[0]))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(excel_data[0]):
        hdr_cells[i].text = str(header)

    for row_data in excel_data[1:]:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)

    marqueur_nom = "Nom==="
    for paragraph in doc.paragraphs:
        if marqueur_nom in paragraph.text:
            # Remplacer le marqueur par le nom et prénom fournis
            texte_remplacement = paragraph.text.replace(marqueur_nom, f"{nom}")
            paragraph.text = texte_remplacement

    marqueur_prenom = "prenom==="
    for paragraph in doc.paragraphs:
        if marqueur_prenom in paragraph.text:
            # Remplacer le marqueur par le  prénom fournis
            texte_remplacement = paragraph.text.replace(marqueur_prenom, f"{prenom}")
            paragraph.text = texte_remplacement

    marqueur_adresse = "adresse==="
    for paragraph in doc.paragraphs:
        if marqueur_adresse in paragraph.text:
            # Remplacer le marqueur par l'adresse fournis
            texte_remplacement = paragraph.text.replace(marqueur_adresse, f"{adresse}")
            paragraph.text = texte_remplacement

    marqueur_tel = "tel==="
    for paragraph in doc.paragraphs:
        if marqueur_tel in paragraph.text:
            # Remplacer le marqueur par l'adresse fournis
            texte_remplacement = paragraph.text.replace(marqueur_tel, f"{tel}")
            paragraph.text = texte_remplacement

    marqueur_tel = "code_postale==="
    for paragraph in doc.paragraphs:
        if marqueur_tel in paragraph.text:
            # Remplacer le marqueur par l'adresse fournis
            texte_remplacement = paragraph.text.replace(marqueur_tel, f"{code_postal}")
            paragraph.text = texte_remplacement

    marqueur_tel = "ville==="
    for paragraph in doc.paragraphs:
        if marqueur_tel in paragraph.text:
            # Remplacer le marqueur par l'adresse fournis
            texte_remplacement = paragraph.text.replace(marqueur_tel, f"{ville}")
            paragraph.text = texte_remplacement



    # Enregistre les modifications dans le document Word
    doc.save("nouveau_document.docx")

    # Imprime le document Word
    printer_name = win32print.GetDefaultPrinter()
    win32api.ShellExecute(0, "print", "nouveau_document.docx", f'"/d:{printer_name}"', ".", 0)

def supprimer():
    global donnees
    workbook = openpyxl.load_workbook('sample_data.xlsx')
    sheet = workbook['feuil1'] if 'feuil1' in workbook.sheetnames else workbook.create_sheet("feuil1")
    # Supprimer la première ligne dans la feuille Excel
    if sheet.max_row >= 1:
        sheet.delete_rows(1, 1)

    # Supprimer la première ligne du tableau de l'interface Tkinter
    if donnees:
        donnees = donnees[1:]  # Supprimer la première ligne de la liste

    workbook.save('sample_data.xlsx')
    tableau_donnee.delete(tableau_donnee.get_children()[0])  # Supprimer la première ligne du tableau dans l'interface

    print('Première ligne supprimée')

def ouvrir_fenetre_profil ():
    fenetre_profil.deiconify()

def enregistrer_fenetre_ajout ():
    workbook = openpyxl.load_workbook('sample_data.xlsx')
    sheet2 = workbook['feuil2'] if 'feuil2' in workbook.sheetnames else workbook.create_sheet("feuil2")
    # Obtenir les valeurs des champs de profil
    nom_excel = entre_nom.get("1.0", "end-1c")
    prenom_excel = entre_prenom.get("1.0", "end-1c")
    adresse_excel = entre_adresse.get("1.0", "end-1c")
    tel_excel = entre_tel.get("1.0", "end-1c")
    code_postal_excel=entre_code_postal.get("1.0", "end-1c")
    ville_excel=entre_ville.get("1.0", "end-1c")

    if len(tel_excel)==10:
        if tel_excel.isdigit() :
            sheet2.cell(row=2, column=11, value=tel_excel)
        else :
            print ('erreur dans le numero de telephone' )
    else :
            print ('erreur dans le numero de telephone' )
    
    for caractere in nom_excel:
        if not caractere.isdigit():
            sheet2.cell(row=2, column=8, value=nom_excel)
        else :
            print ('erreur dans le nom' )
        
    for caractere in prenom_excel:
        if not caractere.isdigit():
            sheet2.cell(row=2, column=9, value=prenom_excel)
        else :
            print ('erreur dans le prenom' )

    if code_postal_excel.isdigit():
        sheet2.cell(row=2, column=12, value=code_postal_excel)
    else:
        print('erreur dans le code postal')

    for caractere in ville_excel:
        if not caractere.isdigit():
            sheet2.cell(row=2, column=13, value=ville_excel)
        else :
            print ('erreur dans la ville ' )

    texte_enregistrement_profil.config(text="enregistrement fini")
    sheet2.cell(row=2, column=10, value=adresse_excel)
    # Sauvegarder les modifications dans le fichier Excel
    workbook.save('sample_data.xlsx')

# creation de la fenetre principale 

fenetre = tk.Tk()
fenetre.geometry("600x400")
fenetre.title("Fenêtre principale")

bouton_imprimer = tk.Button(fenetre, text="Imprimer", command=imprimer, width=15, height=1)
bouton_imprimer.pack()
bouton_imprimer.place(x=1, y=30)

bouton_supprimer = tk.Button(fenetre, text="Supprimer", command=supprimer, width=15, height=1)
bouton_supprimer.pack()
bouton_supprimer.place(x=1, y=60)

bouton_profil = tk.Button(fenetre, text="profil", command=ouvrir_fenetre_profil , width=15, height=1)
bouton_profil.pack()
bouton_profil.place(x=1, y=90)

bouton = tk.Button(fenetre, text="Ajouter un frais", command=ouvrir_fenetre_ajout , width=15, height=1)#bouton pour ajouter un frais
bouton.pack()
bouton.place(x=1, y=1)



#creation de la fenetre profil
fenetre_profil = tk.Toplevel(fenetre)
fenetre_profil.geometry("300x287")
fenetre_profil.title("profil")
fenetre_profil.withdraw()

texte_nom = tk.Label(fenetre_profil, text="nom :")
texte_nom.pack()
texte_nom.place(x=1, y=1)

entre_nom = tk.Text(fenetre_profil, height=1, width=10)
entre_nom.pack()
entre_nom.place(x=80, y=1)

texte_prenom = tk.Label(fenetre_profil, text="prenom :")
texte_prenom.pack()
texte_prenom.place(x=1, y=30)

entre_prenom = tk.Text(fenetre_profil, height=1, width=10)
entre_prenom.pack()
entre_prenom.place(x=80, y=30)

texte_adresse = tk.Label(fenetre_profil, text="adresse :")
texte_adresse.pack()
texte_adresse.place(x=1, y=60)

entre_adresse = tk.Text(fenetre_profil, height=1, width=10)
entre_adresse.pack()
entre_adresse.place(x=80, y=60)

texte_tel = tk.Label(fenetre_profil, text="telephone :")
texte_tel.pack()
texte_tel.place(x=1, y=150)

entre_tel = tk.Text(fenetre_profil, height=1, width=10)
entre_tel.pack()
entre_tel.place(x=80, y=150)

texte_code_postale = tk.Label(fenetre_profil, text="code postal :")
texte_code_postale.pack()
texte_code_postale.place(x=1, y=90)

entre_code_postal = tk.Text(fenetre_profil, height=1, width=10)
entre_code_postal.pack()
entre_code_postal.place(x=80, y=90)

texte_ville = tk.Label(fenetre_profil, text="ville :")
texte_ville.pack()
texte_ville.place(x=1, y=120)

entre_ville = tk.Text(fenetre_profil, height=1, width=10)
entre_ville.pack()
entre_ville.place(x=80, y=120)

bouton_enregistrer_profil = tk.Button(fenetre_profil, text="engistrer", command=enregistrer_fenetre_ajout , width=15, height=1)
bouton_enregistrer_profil.pack()
bouton_enregistrer_profil.place(x=150, y=200)

texte_enregistrement_profil  = tk.Label(fenetre, text="")
texte_enregistrement_profil.place(x=200, y=200)
texte_enregistrement_profil.pack()


#creaation de la fenetre ajout

fenetre_ajout = tk.Toplevel(fenetre)
fenetre_ajout.geometry("300x287")
fenetre_ajout.title("Ajout")
fenetre_ajout.withdraw()

texte_combien = tk.Label(fenetre_ajout, text="montant")
texte_combien.pack()
texte_combien.place(x=1, y=1)

entrée_frais = tk.Text(fenetre_ajout, height=1, width=10)
entrée_frais.pack()
entrée_frais.place(x=80, y=1)

menu_type_frais = tk.Menu(fenetre_ajout)
menu_fichier = tk.Menu(menu_type_frais)
menu_type_frais.add_cascade(label="Type", menu=menu_fichier)
menu_fichier.add_command(label="Hôtel", command=lambda: choisir_type_frais("hotel"))
menu_fichier.add_command(label="Repas", command=lambda: choisir_type_frais("repas"))
menu_fichier.add_command(label="Essence", command=lambda: choisir_type_frais("essence"))
menu_fichier.add_command(label="Péage", command=lambda: choisir_type_frais("peage"))
fenetre_ajout.config(menu=menu_type_frais)

texte_date = tk.Label(fenetre_ajout, text="Date")
texte_date.pack()
texte_date.place(x=1, y=30)

entre_date = tk.Text(fenetre_ajout, height=1, width=10)
entre_date.pack()
entre_date.place(x=80, y=30)

bouton_stocker = tk.Button(fenetre_ajout, text="Stocker le texte", command=stocker_texte)
bouton_stocker.pack()
bouton_stocker.place(x=200, y=200)

texte_enregistrement_ajout = tk.Label(fenetre, text="")
texte_enregistrement_profil.place(x=200 , y=200)
texte_enregistrement_ajout.pack()



# creation du tableau

fenetre['bg'] = '#FFFFFF'
tableau_frame = tk.Frame(fenetre, height=500, width=1800)
tableau_frame.place(x=150, y=1)  # Utilisation de grid pour placer le tableau_frame

tableau_donnee = ttk.Treeview(tableau_frame)
tableau_donnee['columns'] = ('Date', 'Type', 'Montant', 'Indemnisation')

tableau_donnee.column("#0", width=0, stretch='no')
tableau_donnee.column("Date", anchor=tk.CENTER, width=80)
tableau_donnee.column("Type", anchor=tk.CENTER, width=80)
tableau_donnee.column("Montant", anchor=tk.CENTER, width=80)
tableau_donnee.column("Indemnisation", anchor=tk.CENTER, width=80)

tableau_donnee.heading("#0", text="", anchor=tk.CENTER)
tableau_donnee.heading("Date", text="Date", anchor=tk.CENTER)
tableau_donnee.heading("Type", text="Type", anchor=tk.CENTER)
tableau_donnee.heading("Montant", text="Montant", anchor=tk.CENTER)
tableau_donnee.heading("Indemnisation", text="Indemnisation", anchor=tk.CENTER)

for ligne in donnees:
    tableau_donnee.insert(parent='', index='end', values=(ligne[0], ligne[1], ligne[2], ligne[3]))

for row in sheet.iter_rows(values_only=True):
    tableau_donnee.insert(parent='', index='end', values=row)

tableau_donnee.grid(row=0, column=0, sticky='nsew')  # Utilisation de grid pour placer le tableau_donnee dans le frame

# Configuration de la croissance du frame et du treeview avec grid_rowconfigure et grid_columnconfigure
tableau_frame.grid_rowconfigure(0, weight=1)
tableau_frame.grid_columnconfigure(0, weight=1)

fenetre.mainloop()